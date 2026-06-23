from docx import Document
import json

mapping = json.load(open("mapping.json"))
reverse = {v: k for k, v in mapping.items()}

def deanonymize_text(text):
    for code, name in reverse.items():
        text = text.replace(code, name)
    return text

doc = Document("anonymized.docx")

for p in doc.paragraphs:
    p.text = deanonymize_text(p.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            cell.text = deanonymize_text(cell.text)

doc.save("restored.docx")
